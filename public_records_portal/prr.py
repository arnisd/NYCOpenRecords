#!/usr/bin/python
# -*- coding: utf-8 -*-

"""
    public_records_portal.prr
    ~~~~~~~~~~~~~~~~

    Implements functions specific to managing or creating a public records request.

"""

import csv
import traceback
import urllib
from StringIO import StringIO
import os
import json
import subprocess
import bleach
from public_records_portal import cal
from flask import request, render_template, make_response, send_file
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, Inches

import upload_helpers
from RequestPresenter import RequestPresenter
from ResponsePresenter import ResponsePresenter
from db_helpers import find_request, create_request, get_obj, \
    add_staff_participant, remove_staff_participant, update_obj, \
    get_attribute, change_request_status, create_or_return_user, \
    create_subscriber, create_record, create_note, create_QA, \
    create_answer, update_user, id_generator, get_user_by_id
from models import *
from notifications import generate_prr_emails
from public_records_portal import db_helpers
from requires_roles import requires_roles

agency_codes = {"City Commission on Human Rights": "228",
                "Department of Education": "040",
                "Department of Information Technology and Telecommunications": "858",
                "Department of Records and Information Services": "860",
                "Office of the Mayor": "002",
                "Mayor's Office of Contract Services": "002",
                "Mayor's Office of Media and Entertainment": "002",
                "Office of Administrative Trials and Hearings": "820",
                "Office of the Chief Medical Examiner": "816",
                "NYC Emergency Management": "017",
                None: "000"}


def add_public_note(request_id, text):
    app.logger.info("Add Public Note")
    return 1


def nonportal_request(request_body):
    # Handles requests made to nonportal agencies by sending email notifications to the agency user repsonsible for that
    # department and the user creating a notification

    # query with all the fields
    notification_content = {}
    notification_content['department'] = \
        Department.query.filter_by(name=request_body['request_agency'
        ]).first()
    notification_content['department_name'] = notification_content['department'].name
    notification_content['user_id'] = notification_content['department'].primary_contact_id
    notification_content['recipient'] = User.query.get(notification_content['department'].primary_contact_id).email

    # pass back array with both the request title and the request description

    notification_content['text'] = {}
    notification_content['text']['request_summary'] = request_body['request_summary']
    notification_content['text']['request_text'] = request_body['request_text']

    # checks for the contact information left by the requester

    contact_info = {
        'alias': str(request_body['request_first_name'] + ' '
                     + request_body['request_last_name']),
        'request_role': request_body['request_role'],
        'request_organization': request_body['request_organization'],
        'request_email': request_body['request_email'],
        'request_phone': request_body['request_phone'],
        'request_fax': request_body['request_fax'],
        'request_address_street_one': request_body['request_address_street_one'
        ],
        'request_address_street_two': request_body['request_address_street_two'
        ],
        'request_address_city': request_body['request_address_city'],
        'request_address_state': request_body['request_address_state'],
        'request_address_zip': request_body['request_address_zip'],
    }

    notification_content['contact_info'] = contact_info
    generate_prr_emails(request_id=None,
                        notification_type='non-portal-agency_agency',
                        notification_content=notification_content)

    notification_content['recipient'] = request_body['request_email']
    generate_prr_emails(request_id=None,
                        notification_type='non-portal-agency_requester',
                        notification_content=notification_content)
    return 1


### @export "add_resource"

def add_resource(resource, request_body, current_user_id=None):
    notification_content = {}
    fields = request_body
    department_name = \
        Department.query.filter_by(id=Request.query.filter_by(id=fields['request_id'
        ]).first().department_id).first()
    if 'extension' in resource:
        return request_extension(
                fields['request_id'],
                fields.getlist('extend_reason'),
                current_user_id,
                int(fields['days_after']),
                fields['due_date'],
                request_body=request_body,
        )
    if 'note' in resource:
        return add_note(request_id=fields['request_id'],
                        text=fields['note_text'],
                        user_id=current_user_id,
                        privacy=fields['note_privacy'],
                        request_body=request_body)
    if 'letter' in resource:
        return add_letter(request_id=fields['request_id'],
                          text=fields['response_template'],
                          user_id=current_user_id)
    elif 'record' in resource:
        app.logger.info('''

inside add_resource method''')
        if 'record_access' in fields and fields['record_access'] != "":
            if fields['record_privacy'] == 'release_and_public':
                privacy = RecordPrivacy.RELEASED_AND_PUBLIC
            elif fields['record_privacy'] == 'release_and_private':
                privacy = RecordPrivacy.RELEASED_AND_PRIVATE
            else:
                privacy = RecordPrivacy.PRIVATE
            return add_offline_record(fields['request_id'], bleach.clean(fields['record_description']),
                                      bleach.clean(fields['record_access']),
                                      current_user_id, department_name, privacy=privacy)
        elif 'link_url' in fields and fields['link_url'] != "":
            if fields['record_privacy'] == 'release_and_public':
                privacy = RecordPrivacy.RELEASED_AND_PUBLIC
            elif fields['record_privacy'] == 'release_and_private':
                privacy = RecordPrivacy.RELEASED_AND_PRIVATE
            else:
                privacy = RecordPrivacy.PRIVATE
            return add_link(request_id=bleach.clean(fields['request_id']), url=bleach.clean(fields['link_url']),
                            description=bleach.clean(fields['record_description']), user_id=current_user_id,
                            department_name=department_name, privacy=privacy)
        else:
            app.logger.info('''

everything else...''')
            documents = None
            try:
                documents = request.files.getlist('record')
            except:
                app.logger.info('''

No file passed in''')

            try:
                titles = request.form.getlist('title[]')
            except:
                app.logger.info('''No titles passed in for each record''')

            addAsEmailAttachmentList = []

            for index, t in list(enumerate(titles)):
                addAsEmailAttachment = request.form.get('addAsEmailAttachment_' + str(index))
                if addAsEmailAttachment == 'on':
                    addAsEmailAttachmentList.append(addAsEmailAttachment)
                else:
                    addAsEmailAttachmentList.append(None)

                app.logger.info('addAsEmailAttachment_' + str(index) + ":" + str(addAsEmailAttachment))

            if fields['record_privacy'] == 'release_and_public':
                privacy = RecordPrivacy.RELEASED_AND_PUBLIC
            elif fields['record_privacy'] == 'release_and_private':
                privacy = RecordPrivacy.RELEASED_AND_PRIVATE
            else:
                privacy = RecordPrivacy.PRIVATE

            return upload_multiple_records(
                    request_id=fields['request_id'],
                    documents=documents,
                    addAsEmailAttachmentList = addAsEmailAttachmentList,
                    request_body=request_body,
                    description=fields['record_description'],
                    user_id=current_user_id,
                    privacy=privacy,
                    department_name = department_name,
                    titles=titles)
    elif 'qa' in resource:
        return ask_a_question(request_id=fields['request_id'],
                              user_id=current_user_id,
                              question=fields['question_text'],
                              department_name=department_name)
    elif 'owner' in resource:
        (participant_id, new) = \
            add_staff_participant(request_id=fields['request_id'],
                                  email=fields['owner_email'],
                                  reason=fields['owner_reason'])

        # if new:

        if request_body:

            user = User.query.filter_by(email=request_body['owner_email'
            ]).first()
            user_name = user.alias
            notification_content['user_name'] = user_name
            notification_content['user_id'] = get_attribute('user_id',
                                                            obj_id=participant_id, obj_type='Owner')
            notification_content['request_body'] = request_body
            notification_content['request_id'] = request_body['request_id']
            generate_prr_emails(request_id=fields['request_id'],
                                notification_type='helper_added'
                                , notification_content=notification_content)
        else:
            user = User.query.get(current_user_id)
            user_name = user.alias
            notification_content['user_name'] = user_name
            notification_content['request_id'] = request_body['request_id']
            notification_content['user_id'] = get_attribute('user_id',
                                                            obj_id=participant_id, obj_type='Owner')
            generate_prr_emails(request_id=fields['request_id'],
                                notification_type='helper_removed'
                                , notification_content=notification_content)
        return participant_id
    elif 'subscriber' in resource:
        return add_subscriber(request_id=fields['request_id'],
                              email=fields['follow_email'])
    else:
        return False


### @export "update_resource"

def update_resource(resource, request_body):
    fields = request_body
    notification_content = {}
    if 'owner' in resource:
        if 'change_assignee' in fields:

            # own=Owner.query.filter_by(user_id=fields['owner_id']).first()

            user_name = \
                User.query.filter_by(id=Owner.query.filter_by(id=fields['owner_id'
                ]).first().user_id).first().alias
            notification_content['owner_reason'] = request_body['owner_reason']
            notification_content['user_name'] = user_name
            generate_prr_emails(request_id=fields['request_id'],
                                notification_type='Request assigned',
                                notification_content=notification_content)
        elif 'remove_helper' in fields:
            owner = Owner.query.get(request_body['owner_id'])
            user_id = owner.user_id
            user_name = User.query.get(user_id).alias
            notification_content['user_id'] = user_id
            notification_content['user_name'] = user_name
            notification_content['request_body'] = request_body
            notification_content['request_id'] = request_body['request_id']
            generate_prr_emails(request_id=fields['request_id'],
                                notification_type='helper_removed',
                                notification_content=notification_content)
            return remove_staff_participant(owner_id=owner.id,
                                            reason=fields['reason_unassigned'])
        elif 'reason_unassigned' in fields:
            return remove_staff_participant(owner_id=fields['owner_id'])
        else:
            change_request_status(fields['request_id'], 'Rerouted')
            return assign_owner(request_id=fields['request_id'],
                                reason=fields['owner_reason'],
                                email=fields['owner_email'])
    elif 'reopen' in resource:
        request_id = fields['request_id']
        change_request_status(request_id, 'Reopened')
        req = get_obj('Request', request_id)
        try:
            user_id = req.subscribers[0].user.id
            notification_content['user_id'] = user_id
            generate_prr_emails(request_id=request_id, notification_content=notification_content,
                                notification_type='Public Notification Template 10'
                                )
        except IndexError:
            app.logger.error('No Subscribers')
    elif 'acknowledge' in resource:
        change_request_status(fields['request_id'],
                              fields['acknowledge_status'])
        notification_content['additional_information'] = request_body['additional_information']
        notification_content['acknowledge_status'] = request_body['acknowledge_status']
        generate_prr_emails(request_id=fields['request_id'],
                            notification_content=notification_content,
                            notification_type='acknowledgement')
        return fields['request_id']
    elif 'request_text' in resource:
        update_obj(attribute='text', val=fields['request_text'],
                   obj_type='Request', obj_id=fields['request_id'])
    elif 'note_text' in resource:
        update_obj(attribute='text', val=fields['note_text'],
                   obj_type='Note', obj_id=fields['response_id'])
    elif 'note_delete' in resource:

        # Need to store note text somewhere else (or just do delete here as well)
        # Need to store note somewhere else

        remove_obj('Note', int(fields['response_id']))
    elif 'record_delete' in resource:
        remove_obj('Record', int(fields['record_id']))
    else:
        return False


### @export "request_extension"

@requires_roles('Portal Administrator', 'Agency Administrator',
                'Agency FOIL Officer')
def request_extension(
        request_id,
        extension_reasons,
        user_id,
        days_after=None,
        due_date=None,
        request_body=None,
):
    """
    This function allows portal admins, agency admins, and FOIL officers to request additonal time.
    Uses add_resource from prr.py and takes extension date from field retrived from that function.
    Returns note with new date after adding delta.
    """

    req = Request.query.get(request_id)
    req.extension(days_after, due_date)
    user = User.query.get(user_id)
    user_name = user.alias
    text = 'Request extended:'
    notification_content = {}
    notification_content['days_after'] = days_after
    notification_content['additional_information'] = extension_reasons
    notification_content['due_date'] = str(req.due_date).split(' ')[0]

    print request_body

    if request_body is not None:
        generate_prr_emails(request_id=request_id,
                            notification_type='extension',
                            notification_content=notification_content)
    else:
        generate_prr_emails(request_id=request_id,
                            notification_type='extension',
                            notification_content=notification_content)
    for reason in extension_reasons:
        text = text + reason + '</br>'
    add_staff_participant(request_id=request_id, user_id=user_id)
    return add_note(request_id=request_id, text=text, user_id=user_id,
                    extension=True)  # Bypass spam filter because they are logged in.


def add_note(
        request_id,
        text,
        privacy=1,
        request_body=None,
        user_id=None,
        extension=False,
):
    fields = request_body
    notification_content = {}
    if request_body:
        notification_content['user_id'] = user_id
        notification_content['text'] = request_body['note_text']
        notification_content['additional_information'] = request_body['additional_information']
    if text and text != '':
        note_id = create_note(request_id=request_id, text=text,
                              user_id=user_id, privacy=privacy)
        if extension:
            return note_id
        if note_id:
            change_request_status(request_id,
                                  'A response has been added.')

            if user_id:
                add_staff_participant(request_id=request_id,
                                      user_id=user_id)
                notification_content['request_id'] = request_id
                if request_body is not None:
                    generate_prr_emails(request_id=request_id,
                                        notification_type='add_note',
                                        notification_content=notification_content)
                else:
                    generate_prr_emails(request_id=request_id,
                                        notification_type='add_note',
                                        notification_content=notification_content)
            else:
                if request_body is not None:
                    generate_prr_emails(request_id=request_id,
                                        notification_type='add_note',
                                        notification_content=notification_content)
                else:
                    generate_prr_emails(request_id=request_id,
                                        notification_type='add_note',
                                        notification_content=notification_content)
            return note_id
        return False
    return False


### @export "add_pdf"

def add_pdf(
        request_id,
        text,
        user_id=None,
        passed_spam_filter=False,
        privacy=1,
):
    if 'template_' in text:
        template_num = text.split('_')[1]
        template_name = 'standard_response_' + template_num + '.html'
        app.logger.info('''

PDF TEMPLATE:''' + template_name)
        date = datetime.now().strftime('%B %d, %Y')
        req = Request.query.get(request_id)
        department = \
            Department.query.filter_by(id=req.department_id).first()
        appeals_officer = 'APPEALS OFFICER'
        appeals_email = 'APPEALS_EMAIL'
        staff = req.point_person()
        staff_alias = staff.user.alias
        staff_signature = staff.user.staff_signature
        app.logger.info("Subscribers: %s" % str(req.subscribers))
        if req.subscribers != []:
            user_name = req.subscribers[0]
        else:
            user_name = 'First Name Last Name'

        html = make_response(render_template(
                template_name,
                user_name=user_name,
                date=date,
                req=req,
                department=department,
                appeals_officer=appeals_officer,
                appeals_email=appeals_email,
                staff_alias=staff_alias,
                staff_signature=staff_signature,
        ))
        pdf = StringIO()
        pisaStatus = \
            pisa.CreatePDF(StringIO(html.get_data().encode('utf-8')),
                           pdf)
        if not pisaStatus.err:
            resp = pdf.getvalue()
            pdf.close()
            response = make_response(resp)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = \
                'attachment; filename = Response.pdf'
            return response


def add_letter(
        request_id,
        text,
        user_id=None):
    letters_json = open(os.path.join(app.root_path, 'static/json/letters.json'))
    letters_data = json.load(letters_json)

    date = datetime.now().strftime("%b %d, %Y")
    request = Request.query.get(request_id)
    department = Department.query.filter_by(id=request.department_id).first()
    staff = request.point_person()

    document = Document()

    document.add_paragraph("%s\n" % date)

    requester = request.requester()
    if requester:
        user = requester.user
        alias = user.alias
        address1 = user.address1
        address2 = user.address2
        city = user.city
        state = user.state
        zipcode = user.zipcode
        if address1:
            if address2:
                header = "%s\n%s\n%s\n%s, %s, %s\n\n" % (alias, address1, address2, city, state, zipcode)
            else:
                header = "%s\n%s\n%s, %s, %s\n\n" % (alias, address1, city, state, zipcode)
        else:
            header = "INSERT USER'S ADDRESS HERE"
        document.add_paragraph("%s" % header)

    if text == 'acknowledgement_letter':
        subject = letters_data[text]['subject']
        subject = subject % request_id

        content = letters_data[text]['content']
        content = content % ( str(app.config['PUBLIC_APPLICATION_URL'] + 'request/' + request_id), staff.user.alias)

        document.add_paragraph(subject)
        document.add_paragraph(content)

    if text == 'denial_letter':
        subject = letters_data[text]['subject']
        subject = subject % request_id

        content = letters_data[text]['content']
        content = content % (str(app.config['PUBLIC_APPLICATION_URL'] + 'request/' + request_id), staff.user.alias)

        document.add_paragraph(subject)
        document.add_paragraph(content)

        document = generate_denial_page(document)

    if text == 'extension_letter':
        subject = letters_data[text]['subject']
        subject = subject % request_id

        content = letters_data[text]['content']
        content = content % ( str(app.config['PUBLIC_APPLICATION_URL'] + 'request/' + request_id), staff.user.alias)

        document.add_paragraph(subject)
        document.add_paragraph(content)

    if text == 'partial_fulfillment_letter':
        subject = letters_data[text]['subject']
        subject = subject % request_id

        content = letters_data[text]['content']
        content = content % (str(app.config['PUBLIC_APPLICATION_URL'] + 'request/' + request_id), staff.user.alias)

        document.add_paragraph(subject)
        document.add_paragraph(content)

        document = generate_denial_page(document)

    if text == 'fulfillment_letter':
        subject = letters_data[text]['subject']
        subject = subject % request_id

        content = letters_data[text]['content']
        content = content % (str(app.config['PUBLIC_APPLICATION_URL'] + 'request/' + request_id), staff.user.alias)

        document.add_paragraph(subject)
        document.add_paragraph(content)

    if text == 'payment_letter':
        subject = letters_data[text]['subject']
        subject = subject % request_id

        content = letters_data[text]['content']
        content = content % (str(app.config['PUBLIC_APPLICATION_URL'] + 'request/' + request_id), staff.user.alias)

        document.add_paragraph(subject)
        document.add_paragraph(content)


    letter = StringIO()
    document.save(letter)
    length = letter.tell()
    letter.seek(0)
    return send_file(letter, as_attachment=True, attachment_filename='letter.docx')


def generate_denial_page(document):
    document.add_page_break()
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("Your Freedom of Information Law (FOIL) request was denied, in whole or in part, for the reason(s) checked below:")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     The records are not located in the files of this agency.")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the Freedom of Information Law only requires the disclosure of existing agency records.  It does not require the creation, upon request, of new records.")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     The records which you seek already are available to the public through:")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(1)
    run = paragraph.add_run(u"\u25A1     the OpenData Portal www.nyc.gov/")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(1)
    run = paragraph.add_run(u"\u25A1     the Government Publications Portal www.nyc.gov/records/publications")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(1)
    run = paragraph.add_run(u"\u25A1     New York City 311 www.nyc.gov/")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records are specifically exempted from disclosure by state or federal statute;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records, if disclosed would constitute an unwarranted invasion of personal privacy under State law (subdivision two of section eighty-nine of the Public Officers Law;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records if disclosed would impair present or imminent contract awards or collective bargaining negotiations;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records are trade secrets or are submitted to an agency by a commercial enterprise or derived from information obtained from a commercial enterprise and which if disclosed would cause substantial injury to the competitive position of the subject enterprise;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records are compiled for law enforcement purposes and which, if disclosed, would: ")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.75)
    run = paragraph.add_run(u"i. interfere with law enforcement investigations or judicial proceedings; ii. deprive a person of a right to a fair trial or impartial adjudication; iii. identify a confidential source or disclose confidential information relating to a criminal investigation; or iv. reveal criminal investigative techniques or procedures, except routine techniques and procedures")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records if disclosed could endanger the life or safety of any person;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records are inter-agency or intra-agency materials which are not:")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.75)
    run = paragraph.add_run(u"i. statistical or factual tabulations or data; or ii. instructions to staff that affect the public; or iii. final agency policy or determinations; or iv. external audits, including but not limited to audits performed by the comptroller and the federal government;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records are examination questions or answers which are requested prior to the final administration of such questions;")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.left_indent = Inches(.5)
    run = paragraph.add_run(u"\u25A1     the records or portions of records if disclosed, would jeopardize the capacity of an agency or an entity that has shared information with an agency to guarantee the security of its information technology assets, such assets encompassing both electronic information systems and infrastructures.")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("You may appeal the decision to deny access to material that was redacted in part or withheld in entirety  by contacting the agency's FOIL Appeals Officer within 30 days.")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("The contact information is:")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    return document

def upload_multiple_records(request_id, description, user_id, request_body, documents=None, addAsEmailAttachmentList=None, privacy=RecordPrivacy.PRIVATE, department_name=None, titles=None):
    #for document in documents:
    #    upload_record(request_id, titles[titleIndex], user_id, request_body, document, privacy, department_name)


    documents_size = 0
    for document in documents:
        documents_size += len(document.read())
        document.seek(0)
    if documents_size > app.config['MAX_FILE_SIZE']:
        return "File too large"
    return upload_record(request_id, description, user_id, request_body, documents, addAsEmailAttachmentList, privacy, department_name, titles)

### @export "upload_record"
def upload_record(
        request_id,
        description,
        user_id,
        request_body,
        documents=None,
        addAsEmailAttachmentList=None, 
        privacy=RecordPrivacy.PRIVATE,
        department_name = None,
        titles=None
):
    """ Creates a record with upload/download attributes """

    app.logger.info('''

Begins Upload_record method''')
    notification_content = {}
    titleIndex = 0
    try:

        # doc_id is upload_path

        for document in documents:
            (doc_id, filename, error) = \
            upload_helpers.upload_file(document=document, request_id=request_id, privacy=privacy)
            if error == "file_too_large":
                return "File too large"
            elif error == "file_too_small":
                return "File too small"
            elif doc_id == False:
                return "Extension type '%s' is not allowed." % filename
            else:
                # if str(doc_id).isdigit():
                if str(doc_id) == 'VIRUS_FOUND':
                    return 'There was a virus found in the document you uploaded.'
                if doc_id:

                    # record_id = create_record(doc_id = doc_id, request_id = request_id, user_id = user_id, description = description, filename = filename, url = app.config['HOST_URL'] + doc_id)

                    # use file title instead of description if is provided
                    if titles:
                        description = titles[titleIndex]
                        titleIndex = titleIndex + 1

                    record_id = create_record(
                            doc_id=None,
                            request_id=request_id,
                            user_id=user_id,
                            description=description,
                            filename=filename,
                            url=app.config['HOST_URL'] + str(doc_id),
                            privacy=privacy,
                    )
                    change_request_status(request_id,
                                          'A response has been added.')
                    notification_content['user_id'] = user_id
                    notification_content['department_name'] = department_name
                    notification_content['privacy'] = privacy
    except:
        # print sys.exc_info()[0]
        print traceback.format_exc()
        return 'The upload timed out, please try again.'

    if privacy in [RecordPrivacy.RELEASED_AND_PUBLIC, RecordPrivacy.RELEASED_AND_PRIVATE]:
      if len(addAsEmailAttachmentList) > 0:
        notification_content['documents'] = documents
        for index, addAsEmailAttachment in enumerate(addAsEmailAttachmentList):
          if addAsEmailAttachment:
            # attached_file = app.config['UPLOAD_FOLDER'] + '/' + filename
            notification_content['documents'][index] = documents[index]
            notification_content['index'] = index
            generate_prr_emails(request_id=request_id,
                        notification_type='city_response_added',
                        notification_content=notification_content)
          else:
            notification_content['documents'][index] = None


    add_staff_participant(request_id=request_id,
                        user_id=user_id)
    return 1


### @export "add_offline_record"

def add_offline_record(
        request_id,
        description,
        access,
        user_id,
        department_name,
        request_body=None,
        privacy=True,
):
    """ Creates a record with offline attributes """
    notification_content = {}
    notification_content['description'] = description
    notification_content['department_name'] = department_name
    record_id = create_record(request_id=request_id, user_id=user_id,
                              access=access, description=description,
                              privacy=privacy)  # To create an offline record, we need to know the request ID to which it will be added, the user ID for the person adding the record, how it can be accessed, and a description/title of the record.
    if record_id:
        notification_content['department_name'] = department_name
        change_request_status(request_id, 'A response has been added.')
        generate_prr_emails(request_id=request_id,
                            notification_type='city_response_added',
                            notification_content=notification_content)
        add_staff_participant(request_id=request_id, user_id=user_id)
        return record_id
    return False


### @export "add_link"

def add_link(
        request_id,
        url,
        description,
        user_id,
        department_name=None,
        privacy=True,
):
    """ Creates a record with link attributes """
    notification_content = {}
    record_id = create_record(url=url, request_id=request_id,
                              user_id=user_id, description=description,
                              privacy=privacy)
    if record_id:
        change_request_status(request_id, 'A response has been added.')
        notification_content['url'] = url
        notification_content['description'] = description
        notification_content['department_name'] = department_name
        generate_prr_emails(request_id=request_id,
                            notification_type='city_response_added',
                            notification_content=notification_content)
        add_staff_participant(request_id=request_id, user_id=user_id)
        return record_id
    return False


### @export "make_request"

def make_request(
        agency=None,
        summary=None,
        text=None,
        attachment=None,
        attachment_description=None,
        offline_submission_type=None,
        date_received=None,
        first_name=None,
        last_name=None,
        alias=None,
        role=None,
        organization=None,
        email=None,
        phone=None,
        fax=None,
        street_address_one=None,
        street_address_two=None,
        city=None,
        state=None,
        zip=None,
        user_id=None,
):
    """ Make the request. At minimum you need to communicate which record(s) you want, probably with some text."""

    notification_content = {}
    try:
        is_partner_agency = agency_codes[agency]
    except KeyError:
        return (None, True)
    request = find_request(summary)
    if request:  # Same request already exists
        if agency == request.department_name():
            return (request.id, False)
    assigned_to_email = app.config['DEFAULT_OWNER_EMAIL']
    assigned_to_reason = app.config['DEFAULT_OWNER_REASON']
    if agency:
        app.logger.info('''

Agency chosen: %s''' % agency)
        prr_email = db_helpers.get_contact_by_dept(agency)
        if prr_email:
            assigned_to_email = prr_email
            assigned_to_reason = 'PRR Liaison for %s' % agency
    global currentRequestId
    currentRequestId = id_generator.next()
    id = 'FOIL' + '-' + datetime.now().strftime('%Y') + '-' \
         + agency_codes[agency] + '-' + '%05d' % currentRequestId
    req = get_obj('Request', id)
    while req is not None and req.id == id:
        app.logger.info(req.id + ' VS ' + id)
        currentRequestId = id_generator.next()
        id = 'FOIL' + '-' + datetime.now().strftime('%Y') + '-' \
             + agency_codes[agency] + '-' + '%05d' % currentRequestId
        req = get_obj('Request', id)

    request_id = create_request(  # Actually create the Request object
            id=id,
            agency=agency,
            summary=summary,
            text=text,
            user_id=user_id,
            offline_submission_type=offline_submission_type,
            date_received=date_received,
    )

    # Please don't remove call to assign_owner below

    new_owner_id = assign_owner(request_id=request_id,
                                reason=assigned_to_reason,
                                email=assigned_to_email)  # Assign someone to the request
    open_request(request_id)  # Set the status of the incoming request to "Open"
    if email or alias or phone:
        subscriber_user_id = create_or_return_user(
                email=email,
                alias=alias,
                first_name=first_name,
                last_name=last_name,
                phone=phone,
                fax=fax,
                address1=street_address_one,
                address2=street_address_two,
                city=city,
                state=state,
                zipcode=zip,
        )
        (subscriber_id, is_new_subscriber) = \
            create_subscriber(request_id=request_id,
                              user_id=subscriber_user_id)
        notification_content['user_id'] = subscriber_user_id
        notification_content['summary'] = summary
        notification_content['department_name'] = agency
        notification_content['due_date'] = str(Request.query.filter_by(id=id).first().due_date).split(' ')[0]
        if subscriber_id:
            generate_prr_emails(request_id,
                                notification_type='confirmation',
                                notification_content=notification_content)  # Send them an e-mail notification
            generate_prr_emails(request_id,
                                notification_type='confirmation_agency',
                                notification_content=notification_content)
    if attachment:
        upload_record(request_id=request_id,
                      description=attachment_description,
                      user_id=user_id, request_body=None,
                      documents=attachment)
    return (request_id, True)


#
### @export "add_subscriber"

def add_subscriber(request_id, email):
    notification_content = {}
    user_id = create_or_return_user(email=email)
    notification_content['user_id'] = create_or_return_user(email=email)
    (subscriber_id, is_new_subscriber) = \
        create_subscriber(request_id=request_id, user_id=user_id)
    if subscriber_id:
        generate_prr_emails(request_id,
                            notification_type='Request followed',
                            notification_content=notification_content)
        return subscriber_id
    return False


### @export "ask_a_question"

def ask_a_question(
        request_id,
        user_id,
        question,
        department_name=None,
):
    """ City staff can ask a question about a request they are confused about."""

    notification_content = {}
    req = get_obj('Request', request_id)
    qa_id = create_QA(request_id=request_id, question=question,
                      user_id=user_id)
    notification_content['user_id'] = user_id
    notification_content['department_name'] = department_name
    if qa_id:
        change_request_status(request_id, 'Pending')
        requester = req.requester()
        if requester:
            generate_prr_emails(request_id,
                                notification_type='Question asked',
                                notification_content=notification_content)
        add_staff_participant(request_id=request_id, user_id=user_id)
        return qa_id
    return False


### @export "answer_a_question"

def answer_a_question(
        qa_id,
        answer,
        subscriber_id=None,
        passed_spam_filter=False,
):
    """ A requester can answer a question city staff asked them about their request."""

    if not answer or answer == '' or not passed_spam_filter:
        return False
    else:
        request_id = create_answer(qa_id, subscriber_id, answer)

        # We aren't changing the request status if someone's answered a question anymore, but we could change_request_status(request_id, "Pending")

        generate_prr_emails(request_id=request_id,
                            notification_type='Question answered')
        return True


# add email
### @export "open_request"

def open_request(request_id):
    change_request_status(request_id, 'Open')


### @export "assign_owner"

def assign_owner(request_id, reason=None, email=None):
    """ Called any time a new owner is assigned. This will overwrite the current owner."""
    notification_content = {}
    req = get_obj('Request', request_id)
    past_owner_id = None

    # If there is already an owner, unassign them:

    if req.point_person():
        past_owner_id = req.point_person().id
        past_owner = get_obj('Owner', req.point_person().id)
        update_obj(attribute='is_point_person', val=False,
                   obj=past_owner)
    (owner_id, is_new_owner) = \
        add_staff_participant(request_id=request_id, reason=reason,
                              email=email, is_point_person=True)
    if past_owner_id == owner_id:  # Already the current owner, so don't send any e-mails
        return owner_id

    app.logger.info('''

A new owner has been assigned: Owner: %s'''
                    % owner_id)
    new_owner = get_obj('Owner', owner_id)

    # Update the associated department on request

    update_obj(attribute='department_id',
               val=new_owner.user.department_id, obj=req)
    user_id = get_attribute(attribute='user_id', obj_id=owner_id,
                            obj_type='Owner')

    # Send notifications


    if is_new_owner:
        user = User.query.get(user_id)
        notification_content['user_id'] = user_id
        notification_content['user_name'] = user.alias
        notification_content['request_id'] = request_id
        generate_prr_emails(request_id=request_id,
                            notification_type='request_assigned',
                            notification_content=notification_content)
    return owner_id


### @export "get_request_data_chronologically"

def get_request_data_chronologically(req):
    public = False
    if current_user.is_anonymous:
        public = True
    responses = []
    if not req:
        return responses
    for (i, note) in enumerate(req.notes):
        if not note.user_id:
            responses.append(RequestPresenter(note=note, index=i,
                                              public=public, request=req))
    for (i, qa) in enumerate(req.qas):
        responses.append(RequestPresenter(qa=qa, index=i,
                                          public=public, request=req))
    if not responses:
        return responses
    responses.sort(key=lambda x: x.date(), reverse=True)
    return responses


### @export "get_responses_chronologically"

def get_responses_chronologically(req):
    responses = []
    if not req:
        return responses
    for note in req.notes:
        if note.user_id:

            # Ensure private notes only available to appropriate users

            if note.privacy == 2 and (current_user.is_anonymous
                                      or current_user.role not in ['Portal Administrator'
                    , 'Agency Administrator', 'Agency FOIL Officer',
                                                                   'Agency Helpers']):
                pass
            else:
                responses.append(ResponsePresenter(note=note))
    for record in req.records:
        if current_user.is_authenticated:
            responses.append(ResponsePresenter(record=record))
        else:
            if record.privacy == RecordPrivacy.RELEASED_AND_PUBLIC and record.release_date and record.release_date < datetime.now():
                responses.append(ResponsePresenter(record=record))

    if not responses:
        return responses

    responses.sort(key=lambda x: x.date(), reverse=True)
    if 'Closed' in req.status:
        responses[0].set_icon('icon-archive')  # Set most recent note (closed note)'s icon
    return responses


### @export "set_directory_fields"

def set_directory_fields():
    # Set basic user data

    if 'STAFF_URL' in app.config:

        # This gets run at regular internals via db_users.py in order to keep the staff user list up to date. Before users are added/updated, ALL users get reset to 'inactive', and then only the ones in the current CSV are set to active.

        for user in User.query.filter(User.is_staff == True).all():
            update_user(user=user, is_staff=False)
        csvfile = urllib.urlopen(app.config['STAFF_URL'])
        dictreader = csv.DictReader(csvfile, delimiter=',')
        for row in dictreader:
            create_or_return_user(
                    email=row['email'].lower(),
                    alias=row['name'],
                    phone=row['phone number'],
                    department=row['department name'],
                    is_staff=True,
                    role=row['role'],
            )

        # Set liaisons data (who is a PRR liaison for what department)

        if 'LIAISONS_URL' in app.config:
            csvfile = urllib.urlopen(app.config['LIAISONS_URL'])
            dictreader = csv.DictReader(csvfile, delimiter=',')
            for row in dictreader:
                user = create_or_return_user(email=row['PRR liaison'],
                                             contact_for=row['department name'])
                if row['department name'] != '':
                    set_department_contact(row['department name'],
                                           'primary_contact_id', user)
                    if row['PRR backup'] != '':
                        user = \
                            create_or_return_user(email=row['PRR backup'
                            ], backup_for=row['department name'])
                        set_department_contact(row['department name'],
                                               'backup_contact_id', user)
        else:
            app.logger.info('''

 Please update the config variable LIAISONS_URL for where to find department liaison data for your agency.''')
    else:
        app.logger.info('''

 Please update the config variable STAFF_URL for where to find csv data on the users in your agency.''')
        if 'DEFAULT_OWNER_EMAIL' in app.config \
                and 'DEFAULT_OWNER_REASON' in app.config:
            create_or_return_user(email=app.config['DEFAULT_OWNER_EMAIL'
            ].lower(),
                                  alias=app.config['DEFAULT_OWNER_EMAIL'
                                  ],
                                  department=app.config['DEFAULT_OWNER_REASON'
                                  ], is_staff=True)
            app.logger.info('''

 Creating a single user from DEFAULT_OWNER_EMAIL and DEFAULT_OWNER_REASON for now. You may log in with %s'''
                            % app.config['DEFAULT_OWNER_EMAIL'])
        else:
            app.logger.info('''

 Unable to create any users. No one will be able to log in.''')


### @export "set_department_contact"

def set_department_contact(department_name, attribute_name, user_id):
    department = Department.query.filter(Department.name
                                         == department_name).first()
    user_obj = get_user_by_id(user_id)
    user_email = user_obj.email
    app.logger.info('''

Setting %s For %s as %s'''
                    % (attribute_name, department.name, user_email))
    update_obj(attribute=attribute_name, val=user_id,
               obj_type='Department', obj_id=department.id)


# needs additional information
### @export "close_request"

def close_request(
        request_id,
        reasons='',
        user_id=None,
        request_body=None,
):
    req = get_obj('Request', request_id)
    records = Record.query.filter_by(request_id=request_id).all()
    errors={}
    #Goes through all the records associated with a request and checks if they have agency description filled out.
    for rec in records:
        if req.agency_description == None or req.agency_description == '':
            #Check if the agency description is filled out
            app.logger.info("Agency Description is not filled out")
            if rec.access != None:
                #Check if an offline record exists
                app.logger.info("Request contains an offline record!")
                errors['missing_agency_description'] = "You must provide an Agency Description before closing this request since you added offline instructions"
            if (rec.privacy == 0x1) or (rec.privacy == 0x2):
                #Check if there are any documents uploaded which are private
                errors['missing_agency_description_record_privacy'] = "You must provide an Agency Description before closing this request since one or more document is marked as 'Private' or 'Released and Private' "
    if errors:
        return errors
    change_request_status(request_id, 'Closed')
    notification_content = {}
    # Create a note to capture closed information:
    if request_body:
        notification_content['additional_information'] = request_body['additional_information']
    notification_content['explanations'] = []
    notification_content['reasons'] = reasons
    # for reason in reasons:
    #     notification_content['explanations'].append(explain_action(reason, explanation_type='What'))
    notification_content['user_id'] = user_id
    create_note(request_id, reasons, user_id, privacy=1)
    if "Your request under the Freedom of Information Law (FOIL) has been reviewed and the documents you requested have " \
       "been posted on the OpenRecords portal." in notification_content['reasons']:
        notification_type = "closed_fulfilled_in_whole"
    elif "Your request under the Freedom of Information Law (FOIL) has been reviewed and is granted in part and denied " \
         "in part because some of the records or portions of records are disclosable and others are exempt form " \
         "disclosure under FOIL." in notification_content['reasons']:
        notification_type = "closed_fulfilled_in_part"
    elif "Your request under the Freedom of Information Law (FOIL) has been fulfilled and the documents you requested " \
         "have been emailed to you." in notification_content['reasons']:
        notification_type = "closed_by_email"
    elif "Your request under the Freedom of Information Law (FOIL) has been fulfilled and the documents you requested " \
         "have been faxed to you." in notification_content['reasons']:
        notification_type = "closed_by_fax"
    elif "Your request under the Freedom of Information Law (FOIL) has been been fulfilled and the documents you " \
         "requested have been mailed to you. Please allow 7 - 10 days for receipt." in notification_content['reasons']:
        notification_type = "closed_by_mail"
    elif "Your request under the Freedom of Information Law has been fulfilled and the documents you requested are " \
         "available to you for pickup." in notification_content['reasons']:
        notification_type = "closed_by_pickup"
    elif "Your request is for general City information or services. You may search http://nyc.gov/ or call 311 for " \
         "assistance." in notification_content['reasons']:
        notification_type = "refer_311"
    elif "Your request is for data that is available through the City's OpenData site. You may visit " \
         "http://nyc.gov/opendata to find this information" in notification_content['reasons']:
        notification_type = "refer_opendata"
    elif "Your request under the Freedom of Information Law (FOIL) is being closed because this agency does not have " \
         "the records requested. You should direct your request to a different agency." in notification_content['reasons']:
        notification_type= "refer_other_agency"
    elif "Your request under the Freedom of Information Law (FOIL) is being closed because the document or information " \
         "you requested is publicly available." in notification_content['reasons']:
        notification_type = "refer_web_link"
    else:
        # If request is not one of the fulfilled or referred categories, it is assumed to have been denied
        notification_type = "denied"
    if request_body:
        generate_prr_emails(request_id=request_id,
                            notification_type=notification_type,
                            notification_content=notification_content)
    else:
        generate_prr_emails(request_id=request_id,
                            notification_type=notification_type,
                            notification_content=notification_content)
    add_staff_participant(request_id=request_id, user_id=user_id)

    #Update the time of when the agency description should be released to the public to be 10 days from now
    updated_due_date = datetime.now() + timedelta(days=10)
    update_obj(attribute='agency_description_due_date', val=updated_due_date,obj_type='Request', obj_id=req.id)
    return None


def change_privacy_setting(request_id, privacy, field):
    req = get_obj('Request', request_id)
    if (req.description_private==True and privacy==u'True') or (req.title_private==True and privacy==u'True'):
        if req.agency_description == None or req.agency_description == u'':
            return "An Agency Description must be provided if both the description and title are set to private"
    if field == 'title':
    # Set the title to private
        update_obj(attribute='title_private', val=privacy,
                   obj_type='Request', obj_id=req.id)
    elif field == 'description':

        # Set description to private

        req.description_private = privacy
        update_obj(attribute='title_private', val=privacy,
                   obj_type='Request', obj_id=req.id)


def change_record_privacy(record_id, privacy):
    record = get_obj("Record", record_id)
    if privacy == 'release_and_public':
        privacy = RecordPrivacy.RELEASED_AND_PUBLIC
        release_date = cal.addbusdays(datetime.now(), int(app.config['DAYS_TO_POST']))
        update_obj(attribute="release_date", val=None, obj_type="Record", obj_id=record.id)
    elif privacy == 'release_and_private':
        privacy = RecordPrivacy.RELEASED_AND_PRIVATE
        update_obj(attribute="release_date", val=None, obj_type="Record", obj_id=record.id)
    else:
        privacy = RecordPrivacy.PRIVATE
        update_obj(attribute="release_date", val=None, obj_type="Record", obj_id=record.id)
    update_obj(attribute="privacy", val=privacy, obj_type="Record", obj_id=record.id)
    app.logger.info('Syncing privacy changes to %s' % app.config['PUBLIC_SERVER_HOSTNAME'])
    if record.filename and privacy == RecordPrivacy.RELEASED_AND_PUBLIC:
        app.logger.info("Making %s public" % record.filename)
        subprocess.call(["mv", app.config['UPLOAD_PRIVATE_LOCAL_FOLDER'] + "/" + record.filename, app.config['UPLOAD_PUBLIC_LOCAL_FOLDER'] + "/"])
        if app.config['PUBLIC_SERVER_HOSTNAME'] is not None:
            subprocess.call(["rsync", "-avzh", "ssh", app.config['UPLOAD_PUBLIC_LOCAL_FOLDER'] + "/" + record.filename, app.config['PUBLIC_SERVER_USER'] + '@' + app.config['PUBLIC_SERVER_HOSTNAME'] + ':' + app.config['UPLOAD_PUBLIC_REMOTE_FOLDER'] + "/"])
        else:
            subprocess.call(["mv", app.config['UPLOAD_PUBLIC_LOCAL_FOLDER'] + "/" + record.filename, app.config['UPLOAD_PUBLIC_REMOTE_FOLDER'] + "/"])
    elif record.filename and (privacy == RecordPrivacy.RELEASED_AND_PRIVATE or privacy == RecordPrivacy.PRIVATE):
        app.logger.info("Making %s private" % record.filename)
        subprocess.call(["mv", app.config['UPLOAD_PUBLIC_LOCAL_FOLDER'] + "/" + record.filename, app.config['UPLOAD_PRIVATE_LOCAL_FOLDER'] + "/"])
        if app.config['PUBLIC_SERVER_HOSTNAME'] is not None:
            subprocess.call(["rsync", "-avzh", "--delete", "ssh", app.config['UPLOAD_PUBLIC_LOCAL_FOLDER'] + "/" + record.filename, app.config['PUBLIC_SERVER_USER'] + '@' + app.config['PUBLIC_SERVER_HOSTNAME'] + ':' + app.config['UPLOAD_PUBLIC_REMOTE_FOLDER'] + "/"])
        else:
            subprocess.call(["rm", "-rf", app.config['UPLOAD_PUBLIC_REMOTE_FOLDER'] + "/" + record.filename])

    update_obj(attribute="privacy", val=privacy, obj_type="Record", obj_id=record.id)

def edit_agency_description(request_id, agency_description_text):
    #edit the agency description field of the request
    app.logger.info("Modifying agency description of the request")
    update_obj(attribute='agency_description', val=agency_description_text, obj_type='Request', obj_id=request_id)



